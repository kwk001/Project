#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新技术方案3.3章节，添加四大系统各页面详细说明
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading

def add_paragraph(doc, text, bold=False, indent=True):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    return para

def add_list_item(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    return para

# 页面详细说明数据（从功能清单提取）
system_pages = {
    "海关信息系统": {
        "智能卡口控制": [
            {"page": "设备管理", "func": "增/删/改/查", "p": "P0", "desc": "维护地磅、车牌识别设备、道闸、显示屏等设备信息，设备状态实时监控，设备故障自动报警，设备维护记录"},
            {"page": "车辆识别", "func": "查/识别", "p": "P0", "desc": "车牌自动识别、车辆信息比对、异常车辆预警，车牌识别准确率>99%，与预约信息比对，黑名单车辆自动拦截"},
            {"page": "地磅称重", "func": "查/称重", "p": "P0", "desc": "称重数据采集、超重预警、称重记录查询、数据防篡改，称重数据自动记录，超重自动预警，数据防篡改机制"},
            {"page": "金关对接", "func": "查/对接", "p": "P0", "desc": "与海关金关二期系统对接，申报数据比对、运输工具校验、风险预警、指令触发"},
            {"page": "安检对接", "func": "查/对接", "p": "P0", "desc": "与机场、海关安检系统对接，安检申报、安检结果接收、安检状态同步"}
        ],
        "查验预警处置": [
            {"page": "查验筛选", "func": "查/筛选", "p": "P0", "desc": "AI风险预警模型、查验对象智能筛选、查验优先级排序，基于历史数据和风险模型智能筛选"},
            {"page": "指令下发", "func": "增/删/改/查", "p": "P0", "desc": "查验指令生成、指令推送至PDA、查验项目提示，指令实时推送，查验项目自动提示，查验人员自动分配"},
            {"page": "结果录入", "func": "增/删/改/查", "p": "P0", "desc": "查验结果录入、拍照留痕、结果自动同步，PDA端结果录入，查验过程拍照，结果实时同步至海关系统"},
            {"page": "异常处置", "func": "增/删/改/查", "p": "P0", "desc": "异常登记、处置流程、处置审核、处置记录，异常自动识别，标准化处置流程，处置审核确认"}
        ],
        "货物预报管理": [
            {"page": "预报录入", "func": "增/改/查", "p": "P0", "desc": "录入货物基础信息、航空运输信息、收发货人信息，支持手工录入和Excel批量导入，货物信息完整性校验"},
            {"page": "预报审核", "func": "审核", "p": "P0", "desc": "对货物预报信息进行审核，审核通过后推送至海关系统"},
            {"page": "预录入回执", "func": "查/接收", "p": "P0", "desc": "接收海关系统预录入回执，回执解析与展示"}
        ],
        "跨境电商": [
            {"page": "9610业务管理", "func": "增/删/改/查", "p": "P0", "desc": "跨境电商B2C出口业务管理，订单申报、清单申报、汇总申报"},
            {"page": "9710业务管理", "func": "增/删/改/查", "p": "P0", "desc": "跨境电商B2B直接出口业务管理，订单申报、报关单申报"},
            {"page": "9810业务管理", "func": "增/删/改/查", "p": "P0", "desc": "跨境电商出口海外仓业务管理，海外仓备案、出口申报"},
            {"page": "跨境电商海关对接", "func": "查/对接", "p": "P0", "desc": "与海关跨境电商系统对接，数据报送、状态查询、回执接收"}
        ],
        "海关业务": [
            {"page": "海关申报", "func": "增/删/改/查", "p": "P0", "desc": "向海关进行货物申报，包括报关单、清单等单证申报"},
            {"page": "税费计算", "func": "查/算", "p": "P0", "desc": "根据海关税则计算应缴税费"}
        ],
        "可视化展示": [
            {"page": "业务数据大屏", "func": "查/展示", "p": "P1", "desc": "大屏展示业务数据统计和分析"},
            {"page": "监管数据大屏", "func": "查/展示", "p": "P1", "desc": "大屏展示海关监管相关数据"},
            {"page": "卡口运行状况大屏", "func": "查/展示", "p": "P1", "desc": "大屏展示卡口实时运行状态"}
        ]
    },
    "航空货运信息管理系统": {
        "基础资料管理-航班信息": [
            {"page": "航班信息录入", "func": "增/改/查", "p": "P0", "desc": "录入航班基础信息（航班号、起降时间、机型等），为货物跟踪、打板作业提供数据支撑"},
            {"page": "航班信息审核", "func": "审核", "p": "P0", "desc": "审核录入的航班信息，确保数据准确性"},
            {"page": "航班状态变更", "func": "改", "p": "P0", "desc": "实时更新航班动态状态（计划、起飞、到达、延误、取消），触发关联业务"}
        ],
        "基础资料管理-客户信息": [
            {"page": "客户信息录入", "func": "增/改/查", "p": "P0", "desc": "录入客户基础信息（企业信息、联系人、地址等），建立客户档案"},
            {"page": "信用额度控制", "func": "改", "p": "P1", "desc": "管理客户信用额度，控制业务风险，超额度自动预警"}
        ],
        "基础资料管理-货物信息": [
            {"page": "货物信息录入", "func": "增/改/查", "p": "P0", "desc": "录入货物基础信息（品名、规格、重量、体积、海关编码等）"},
            {"page": "货物类型维护", "func": "类型维护", "p": "P1", "desc": "维护货物类型分类，便于业务统计和分析"}
        ],
        "基础资料管理-仓库设置": [
            {"page": "仓库信息维护", "func": "增/改/查", "p": "P0", "desc": "维护仓库基础信息（仓库编码、名称、地址、库容等）"},
            {"page": "库位管理", "func": "增/删/改/查", "p": "P0", "desc": "管理仓库库位信息，支持库位状态（空闲/占用/禁用）管理"},
            {"page": "容器管理", "func": "增/删/改/查", "p": "P1", "desc": "管理ULD容器（集装器）信息，包括类型、规格、状态"}
        ],
        "订单管理": [
            {"page": "订单创建", "func": "增", "p": "P0", "desc": "创建物流订单，录入货物信息、收发货人信息、服务要求"},
            {"page": "订单拆分", "func": "增/删/改/查", "p": "P0", "desc": "将大订单拆分为多个子订单，便于分别处理和跟踪"},
            {"page": "订单变更", "func": "改", "p": "P0", "desc": "修改已生效订单的信息，记录变更历史"},
            {"page": "订单可视化", "func": "可视化", "p": "P1", "desc": "可视化展示订单全生命周期状态和进度"}
        ],
        "运单管理": [
            {"page": "运单录入", "func": "增/改/查", "p": "P0", "desc": "录入航空运单信息（主单/分单、件数、重量、目的地等）"},
            {"page": "运单跟踪", "func": "查", "p": "P0", "desc": "实时跟踪运单状态和货物位置，提供全程可视化"},
            {"page": "运单变更", "func": "改", "p": "P0", "desc": "修改已录入运单的信息，支持改单、退单操作"}
        ],
        "运输调度": [
            {"page": "调度计划", "func": "增/删/改/查", "p": "P0", "desc": "制定运输调度计划，安排车辆、司机、路线、时间"},
            {"page": "车辆配载", "func": "增/删/改/查", "p": "P0", "desc": "优化车辆货物配载，提高装载率"},
            {"page": "路线规划", "func": "增/删/改/查", "p": "P1", "desc": "规划最优运输路线，考虑距离、时间、成本、限行等因素"}
        ],
        "车辆调度": [
            {"page": "预约申请", "func": "增/删/改/查", "p": "P0", "desc": "车辆到场前提交预约申请（车牌、货物、预计到达时间）"},
            {"page": "预约审核", "func": "审核", "p": "P0", "desc": "审核车辆预约申请，确认入场资格"},
            {"page": "智能排队", "func": "查/排队", "p": "P0", "desc": "智能排序调度到场车辆，优化进场顺序"},
            {"page": "优先级计算", "func": "查/计算", "p": "P0", "desc": "基于货物优先级、预约时间、车辆类型计算进场优先级"},
            {"page": "调度指令", "func": "查/指令", "p": "P0", "desc": "向车辆下发调度指令（进场、等候、月台分配等）"}
        ],
        "运输执行": [
            {"page": "运输任务管理", "func": "增/删/改/查", "p": "P0", "desc": "管理运输任务的分配、执行、完成全过程"},
            {"page": "交接单管理", "func": "增/删/改/查", "p": "P0", "desc": "管理货物交接单据，记录交接时间、地点、人员"},
            {"page": "签收管理", "func": "增/删/改/查", "p": "P0", "desc": "管理货物签收过程，支持电子签名、拍照留痕"},
            {"page": "运单图片管理", "func": "增/删/改/查", "p": "P1", "desc": "管理运单相关图片（货物照片、签收照片等）"},
            {"page": "回单管理", "func": "增/删/改/查", "p": "P0", "desc": "管理运输回单，确认运输任务完成"}
        ],
        "月台管理": [
            {"page": "月台自动分配", "func": "查/分配", "p": "P0", "desc": "基于车辆类型、货物类型、作业时间智能分配月台资源"},
            {"page": "冲突检测", "func": "查/检测", "p": "P0", "desc": "检测月台分配冲突（时间冲突、资源冲突），自动预警"},
            {"page": "任务分配", "func": "查/分配", "p": "P0", "desc": "分配装卸任务给作业人员，支持任务派单和抢单模式"},
            {"page": "进度跟踪", "func": "查/跟踪", "p": "P0", "desc": "实时跟踪装卸作业进度，预估完成时间"},
            {"page": "信息采集", "func": "增/采集", "p": "P0", "desc": "采集装卸作业相关信息（实际时间、人员、异常情况）"}
        ],
        "监管仓管理-入库": [
            {"page": "收货预约", "func": "增/删/改/查", "p": "P0", "desc": "提前预约收货时间、库位，安排收货准备"},
            {"page": "到货登记", "func": "增/登记", "p": "P0", "desc": "登记到货信息（车牌、货物、件数、外包装状态）"},
            {"page": "查验配合", "func": "查/配合", "p": "P0", "desc": "配合海关查验，提供货物信息和查验场地"},
            {"page": "上架作业", "func": "增/删/改/查", "p": "P0", "desc": "执行货物上架作业，扫描库位码确认"},
            {"page": "上架单", "func": "增/删/改/查", "p": "P1", "desc": "生成和管理上架作业单"},
            {"page": "收货记录", "func": "增/删/改/查", "p": "P1", "desc": "记录收货详细信息，支持追溯查询"},
            {"page": "取消收货", "func": "删", "p": "P1", "desc": "取消已登记的收货记录，释放资源"}
        ],
        "监管仓管理-出库": [
            {"page": "出库申请", "func": "增/申请", "p": "P0", "desc": "提交出库申请，指定出库货物和时间"},
            {"page": "拣货作业", "func": "增/作业", "p": "P0", "desc": "执行拣货作业，按出库单拣选货物"},
            {"page": "复核校验", "func": "查/校验", "p": "P0", "desc": "复核出库货物，核对品名、件数、重量"},
            {"page": "货物交接", "func": "增/交接", "p": "P0", "desc": "与承运人进行货物交接，确认交接信息"},
            {"page": "发货单", "func": "增/删/改/查", "p": "P1", "desc": "生成和管理发货单据"},
            {"page": "上机确认", "func": "确认", "p": "P0", "desc": "确认货物已装机，完成出库流程"}
        ],
        "货包机集货跟踪-空运": [
            {"page": "航班动态监控", "func": "动态监控", "p": "P0", "desc": "实时监控航班动态（计划、起飞、到达、延误）"},
            {"page": "货物装机确认", "func": "装机确认", "p": "P0", "desc": "确认货物已完成装机，更新运单状态"},
            {"page": "在途跟踪", "func": "查/跟踪", "p": "P0", "desc": "跟踪空运货物在途状态"}
        ],
        "货包机集货跟踪-陆运": [
            {"page": "GPS定位跟踪", "func": "GPS跟踪", "p": "P0", "desc": "通过GPS实时跟踪运输车辆位置"},
            {"page": "轨迹回放", "func": "查/回放", "p": "P1", "desc": "回放车辆历史行驶轨迹"},
            {"page": "在途监控", "func": "查/监控", "p": "P0", "desc": "监控陆运货物在途状态，异常自动预警"}
        ],
        "分拣线控制管理": [
            {"page": "分拣线仓储管理", "func": "仓储管理", "p": "P0", "desc": "管理分拣线缓存区货物存储"},
            {"page": "分拣线全流程跟踪", "func": "流程跟踪", "p": "P0", "desc": "跟踪货物在分拣线上的全流程"},
            {"page": "分拣控制", "func": "增/删/改/查", "p": "P0", "desc": "控制分拣设备执行分拣操作"},
            {"page": "流程调度", "func": "增/删/改/查", "p": "P0", "desc": "调度分拣流程，优化分拣效率"},
            {"page": "打板位资源调度", "func": "资源调度", "p": "P0", "desc": "调度打板位资源，安排打板作业"},
            {"page": "PLC对接与数据通讯", "func": "PLC对接", "p": "P0", "desc": "与分拣线PLC设备对接，实现数据通讯和控制"}
        ]
    },
    "园区运行中心": {
        "态势监控": [
            {"page": "园区态势监控", "func": "查/监控", "p": "P1", "desc": "大屏展示园区整体运行态势"},
            {"page": "交通态势监控", "func": "查/监控", "p": "P1", "desc": "监控园区交通流量、拥堵情况"},
            {"page": "消防态势监控", "func": "查/监控", "p": "P1", "desc": "监控消防设施状态、火警预警"},
            {"page": "资产态势监控", "func": "查/监控", "p": "P2", "desc": "监控园区资产状态和利用率"},
            {"page": "货物态势监控", "func": "查/监控", "p": "P1", "desc": "监控园区货物流转态势"},
            {"page": "车辆态势监控", "func": "查/监控", "p": "P1", "desc": "监控园区车辆分布和运行状态"},
            {"page": "仓储态势监控", "func": "查/监控", "p": "P1", "desc": "监控仓库利用率、库存状态"},
            {"page": "月台态势监控", "func": "查/监控", "p": "P1", "desc": "监控月台占用率和作业状态"},
            {"page": "客服态势监控", "func": "查/监控", "p": "P2", "desc": "监控客服工作量和服务质量"},
            {"page": "安防态势监控", "func": "查/监控", "p": "P1", "desc": "监控园区安防状态和异常事件"},
            {"page": "能耗态势监控", "func": "查/监控", "p": "P2", "desc": "监控园区能耗使用情况"},
            {"page": "综合态势展示", "func": "态势展示", "p": "P1", "desc": "综合展示园区多维度态势"}
        ],
        "园区管理": [
            {"page": "车辆无感通行", "func": "查/通行", "p": "P1", "desc": "实现园区车辆无感通行，自动识别放行"},
            {"page": "车位管理与引导", "func": "增/删/改/查", "p": "P1", "desc": "管理园区停车位，提供车位引导和预约"},
            {"page": "巡更巡检管理", "func": "巡检管理", "p": "P2", "desc": "管理安保巡更巡检工作"},
            {"page": "设备知识库", "func": "增/删/改/查", "p": "P2", "desc": "构建设备故障知识库，辅助维修"},
            {"page": "停车缴费", "func": "查/缴费", "p": "P1", "desc": "支持停车费用查询和缴纳"},
            {"page": "紧急放行", "func": "查/放行", "p": "P1", "desc": "紧急情况下的快速放行通道"}
        ],
        "数字平台": [
            {"page": "物联网平台", "func": "查/平台", "p": "P1", "desc": "统一管理园区IoT设备，数据采集和控制"},
            {"page": "数据治理平台", "func": "增/删/改/查", "p": "P1", "desc": "实现数据标准化、清洗、质量管理"},
            {"page": "视频转码平台", "func": "查/转码", "p": "P2", "desc": "监控视频转码和分发"},
            {"page": "统一认证平台", "func": "查/认证", "p": "P1", "desc": "提供统一的身份认证服务"}
        ]
    },
    "航空物流公共信息平台": {
        "移动端": [
            {"page": "司机APP", "func": "司机APP", "p": "P0", "desc": "司机使用的移动APP，支持预约、导航、签到、查询等功能"},
            {"page": "作业PDA", "func": "增/删/改/查", "p": "P0", "desc": "仓库作业人员使用的PDA终端，支持扫码、入库、出库、盘点"},
            {"page": "查验PDA", "func": "增/删/改/查", "p": "P0", "desc": "海关查验人员使用的PDA终端，支持指令接收、结果录入、拍照"},
            {"page": "管理小程序", "func": "查/管理", "p": "P1", "desc": "管理人员使用的小程序，支持数据查询、审批、监控"},
            {"page": "客户小程序", "func": "查/服务", "p": "P1", "desc": "客户使用的小程序，支持订单查询、货物跟踪、在线下单"}
        ],
        "公共信息服务": [
            {"page": "航班动态查询", "func": "动态查询", "p": "P1", "desc": "提供航班动态信息查询服务"},
            {"page": "货物跟踪查询", "func": "跟踪查询", "p": "P1", "desc": "提供货物全程跟踪查询服务"},
            {"page": "物流资讯发布", "func": "资讯发布", "p": "P2", "desc": "发布物流行业资讯和公告"},
            {"page": "在线服务申请", "func": "服务申请", "p": "P1", "desc": "支持在线提交服务申请"},
            {"page": "物流跟踪服务", "func": "查/跟踪", "p": "P1", "desc": "提供物流跟踪API服务"},
            {"page": "航班动态服务", "func": "查/订阅", "p": "P1", "desc": "提供航班动态订阅服务"},
            {"page": "在线下单服务", "func": "增/下单", "p": "P1", "desc": "支持客户在线下单"}
        ],
        "客服业务": [
            {"page": "在线客服", "func": "查/服务", "p": "P1", "desc": "提供在线客服咨询服务"},
            {"page": "投诉建议管理", "func": "投诉管理", "p": "P1", "desc": "管理客户投诉和建议"},
            {"page": "工单管理", "func": "工单管理", "p": "P1", "desc": "管理客服工单，跟踪处理进度"},
            {"page": "客服知识库", "func": "增/删/改/查", "p": "P2", "desc": "构建客服知识库，提升服务效率"}
        ],
        "系统管理": [
            {"page": "用户权限管理", "func": "增/删/改/查", "p": "P0", "desc": "管理系统用户和权限分配"},
            {"page": "系统参数配置", "func": "增/删/改/查", "p": "P1", "desc": "配置系统运行参数"},
            {"page": "用户登录", "func": "查/登录", "p": "P0", "desc": "提供用户登录认证功能"}
        ],
        "运维服务": [
            {"page": "系统监控", "func": "查/监控", "p": "P1", "desc": "监控系统运行状态和性能指标"},
            {"page": "日志管理", "func": "增/删/改/查", "p": "P1", "desc": "管理系统操作日志和运行日志"},
            {"page": "备份恢复", "func": "增/删/改/查", "p": "P1", "desc": "管理系统数据备份和恢复"}
        ]
    }
}

# 打开文档
doc = Document('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')

# 找到3.3章节后插入详细页面说明
# 我们需要在3.3.4功能架构设计之后的内容之前插入
print("正在查找插入位置...")

# 重新构建3.3章节的内容
# 由于python-docx不便于精确定位和修改，我们采用重新生成3.3章节的方式

print("生成新的3.3章节内容...")

# 创建新文档，包含完整内容
new_doc = Document()

# 复制原文档样式
style = new_doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# 添加3.3功能架构设计章节
add_heading(new_doc, '3.3 功能架构设计', 2)
p = new_doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run('系统功能架构包含四大子系统，共计171个功能点，以下对各系统功能模块和页面进行详细说明：')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

# ===== 航空货运信息管理系统 =====
p = new_doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run('【航空货运信息管理系统】104个功能点')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.bold = True
p = new_doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('航空货运信息管理系统是平台的核心业务系统，涵盖从订单创建到货物交付的全流程管理。')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

for module, pages in system_pages["航空货运信息管理系统"].items():
    p = new_doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'◆ {module}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = True

    # 创建表格展示页面详情
    table = new_doc.add_table(rows=len(pages)+1, cols=4)
    table.style = 'Table Grid'

    # 表头
    headers = ['页面名称', '功能操作', '优先级', '功能说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'E7E6E6')

    # 数据行
    for i, page in enumerate(pages, 1):
        table.rows[i].cells[0].text = page['page']
        table.rows[i].cells[1].text = page['func']
        table.rows[i].cells[2].text = page['p']
        table.rows[i].cells[3].text = page['desc'][:60] + '...' if len(page['desc']) > 60 else page['desc']

new_doc.add_page_break()

# ===== 海关信息系统 =====
p = new_doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run('【海关信息系统】23个功能点')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.bold = True

p = new_doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('海关信息系统实现与海关金关二期系统的深度对接，支持智能卡口控制、查验预警、货物预报等海关监管业务。')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

for module, pages in system_pages["海关信息系统"].items():
    p = new_doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'◆ {module}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = True

    table = new_doc.add_table(rows=len(pages)+1, cols=4)
    table.style = 'Table Grid'

    headers = ['页面名称', '功能操作', '优先级', '功能说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'E7E6E6')

    for i, page in enumerate(pages, 1):
        table.rows[i].cells[0].text = page['page']
        table.rows[i].cells[1].text = page['func']
        table.rows[i].cells[2].text = page['p']
        table.rows[i].cells[3].text = page['desc'][:60] + '...' if len(page['desc']) > 60 else page['desc']

new_doc.add_page_break()

# ===== 园区运行中心 =====
p = new_doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run('【园区运行中心】22个功能点')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.bold = True

p = new_doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('园区运行中心实现园区态势感知和运营管理，支持多维度态势监控和园区综合管理。')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

for module, pages in system_pages["园区运行中心"].items():
    p = new_doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'◆ {module}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = True

    table = new_doc.add_table(rows=len(pages)+1, cols=4)
    table.style = 'Table Grid'

    headers = ['页面名称', '功能操作', '优先级', '功能说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'E7E6E6')

    for i, page in enumerate(pages, 1):
        table.rows[i].cells[0].text = page['page']
        table.rows[i].cells[1].text = page['func']
        table.rows[i].cells[2].text = page['p']
        table.rows[i].cells[3].text = page['desc'][:60] + '...' if len(page['desc']) > 60 else page['desc']

new_doc.add_page_break()

# ===== 航空物流公共信息平台 =====
p = new_doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run('【航空物流公共信息平台】22个功能点')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.bold = True

p = new_doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('航空物流公共信息平台提供多端接入能力和公共信息服务，支持司机APP、作业PDA、管理小程序、客户小程序等多种终端。')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

for module, pages in system_pages["航空物流公共信息平台"].items():
    p = new_doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'◆ {module}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = True

    table = new_doc.add_table(rows=len(pages)+1, cols=4)
    table.style = 'Table Grid'

    headers = ['页面名称', '功能操作', '优先级', '功能说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'E7E6E6')

    for i, page in enumerate(pages, 1):
        table.rows[i].cells[0].text = page['page']
        table.rows[i].cells[1].text = page['func']
        table.rows[i].cells[2].text = page['p']
        table.rows[i].cells[3].text = page['desc'][:60] + '...' if len(page['desc']) > 60 else page['desc']

# 保存新文档
new_doc.save('3.3章节_四大系统页面详细说明.docx')
print("新文档已生成：3.3章节_四大系统页面详细说明.docx")
print("此文档包含完整的3.3功能架构设计章节，可与原技术方案合并")
