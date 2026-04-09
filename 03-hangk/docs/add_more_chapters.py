#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
补充技术方案第4-14章详细内容
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading

def add_paragraph(doc, text, bold=False, indent=True):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    return para

def add_list_item(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    return para

# 打开已有文档
doc = Document('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')

# ========== 4 数据结构设计 ==========
print("生成第4章 数据结构设计...")
add_heading(doc, '4 数据结构设计', 1)

add_heading(doc, '4.1 逻辑结构设计', 2)
add_paragraph(doc, '逻辑结构设计采用实体-关系（ER）模型，主要实体包括：')
add_paragraph(doc, '基础数据实体：')
add_list_item(doc, '航班信息实体：航班号、航班日期、起降机场、计划/实际起降时间、航班状态、机型等属性')
add_list_item(doc, '客户信息实体：客户编码、客户名称、客户类型、信用额度、联系人、地址等属性')
add_list_item(doc, '货物信息实体：货物名称、货物类型、规格型号、计量单位、海关编码等属性')
add_list_item(doc, '仓库信息实体：仓库编码、仓库名称、仓库类型、库位信息、库容等属性')
add_list_item(doc, '设备资源实体：设备编码、设备名称、设备类型、状态、维护记录等属性')

add_paragraph(doc, '业务数据实体：')
add_list_item(doc, '订单实体：订单编号、客户信息、货物信息、航班信息、订单状态、金额等属性')
add_list_item(doc, '运单实体：运单号、主运单号、订单关联、货物描述、件数、重量、目的地等属性')
add_list_item(doc, '运输任务实体：任务编号、车辆信息、司机信息、路线、计划/实际时间、状态等属性')
add_list_item(doc, '仓储作业单实体：作业单号、作业类型、库位信息、货物信息、操作人员等属性')
add_list_item(doc, '计费账单实体：账单编号、客户信息、费用科目、金额、结算状态等属性')

add_paragraph(doc, '海关数据实体：')
add_list_item(doc, '卡口记录实体：车牌号、入场/出场时间、称重数据、海关状态、安检状态等属性')
add_list_item(doc, '查验任务实体：任务编号、运单关联、风险等级、查验人员、结果、异常标记等属性')
add_list_item(doc, '预报数据实体：预报编号、货物信息、海关编码、申报状态、回执信息等属性')

add_heading(doc, '4.2 物理结构设计', 2)
add_paragraph(doc, '物理数据库设计遵循以下原则：')
add_list_item(doc, '分库分表：按业务域垂直分库（货运库、海关库、园区库、公共库），大表按时间或ID水平分表')
add_list_item(doc, '索引优化：主键索引、外键索引、查询条件字段索引、联合索引')
add_list_item(doc, '数据分区：历史数据按月份分区，便于归档清理和查询优化')
add_list_item(doc, '读写分离：查询操作走从库，写入操作走主库，主从同步延迟控制在1秒内')
add_list_item(doc, '数据归档：超过1年的历史数据归档至对象存储，在线库只保留热数据')

add_heading(doc, '4.3 数据库表结构', 2)
add_paragraph(doc, '核心数据表结构设计如下：', bold=True)

# 航班信息表
add_paragraph(doc, '表1：t_flight_info（航班信息表）', bold=True)
flight_table = doc.add_table(rows=14, cols=4)
flight_table.style = 'Table Grid'
headers = ['字段名', '数据类型', '说明', '约束']
for i, h in enumerate(headers):
    cell = flight_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

flight_fields = [
    ['flight_id', 'VARCHAR(32)', '航班ID', 'PK'],
    ['flight_no', 'VARCHAR(20)', '航班号', 'IDX'],
    ['flight_date', 'DATE', '航班日期', 'IDX'],
    ['departure', 'VARCHAR(50)', '起飞机场', ''],
    ['destination', 'VARCHAR(50)', '目的机场', ''],
    ['planned_departure', 'DATETIME', '计划起飞时间', ''],
    ['planned_arrival', 'DATETIME', '计划到达时间', ''],
    ['actual_departure', 'DATETIME', '实际起飞时间', ''],
    ['actual_arrival', 'DATETIME', '实际到达时间', ''],
    ['flight_status', 'TINYINT', '航班状态', ''],
    ['aircraft_type', 'VARCHAR(20)', '机型', ''],
    ['create_time', 'DATETIME', '创建时间', ''],
    ['update_time', 'DATETIME', '更新时间', '']
]
for i, row_data in enumerate(flight_fields, 1):
    for j, val in enumerate(row_data):
        flight_table.rows[i].cells[j].text = val

add_paragraph(doc, '索引设计：')
add_list_item(doc, 'PRIMARY KEY (flight_id)')
add_list_item(doc, 'INDEX idx_flight_no (flight_no)')
add_list_item(doc, 'INDEX idx_flight_date (flight_date)')
add_list_item(doc, 'INDEX idx_status_date (flight_status, flight_date)')

# 客户信息表
add_paragraph(doc, '表2：t_customer_info（客户信息表）', bold=True)
customer_table = doc.add_table(rows=11, cols=4)
customer_table.style = 'Table Grid'
for i, h in enumerate(headers):
    cell = customer_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

customer_fields = [
    ['customer_id', 'VARCHAR(32)', '客户ID', 'PK'],
    ['customer_code', 'VARCHAR(20)', '客户编码', 'UK'],
    ['customer_name', 'VARCHAR(100)', '客户名称', 'IDX'],
    ['customer_type', 'TINYINT', '客户类型', ''],
    ['credit_limit', 'DECIMAL(18,2)', '信用额度', ''],
    ['contact_name', 'VARCHAR(50)', '联系人', ''],
    ['contact_phone', 'VARCHAR(20)', '联系电话', ''],
    ['address', 'VARCHAR(200)', '地址', ''],
    ['status', 'TINYINT', '状态', 'IDX'],
    ['create_time', 'DATETIME', '创建时间', '']
]
for i, row_data in enumerate(customer_fields, 1):
    for j, val in enumerate(row_data):
        customer_table.rows[i].cells[j].text = val

# 订单信息表
add_paragraph(doc, '表3：t_order_info（订单信息表）', bold=True)
order_table = doc.add_table(rows=17, cols=4)
order_table.style = 'Table Grid'
for i, h in enumerate(headers):
    cell = order_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

order_fields = [
    ['order_id', 'VARCHAR(32)', '订单ID', 'PK'],
    ['order_no', 'VARCHAR(30)', '订单编号', 'UK'],
    ['customer_id', 'VARCHAR(32)', '客户ID', 'FK,IDX'],
    ['order_type', 'TINYINT', '订单类型', 'IDX'],
    ['cargo_name', 'VARCHAR(200)', '货物名称', 'IDX'],
    ['cargo_weight', 'DECIMAL(10,3)', '货物重量(kg)', ''],
    ['cargo_volume', 'DECIMAL(10,3)', '货物体积(m³)', ''],
    ['cargo_quantity', 'INT', '货物件数', ''],
    ['sender_id', 'VARCHAR(32)', '发货人ID', 'FK'],
    ['receiver_id', 'VARCHAR(32)', '收货人ID', 'FK'],
    ['flight_id', 'VARCHAR(32)', '航班ID', 'FK,IDX'],
    ['order_status', 'TINYINT', '订单状态', 'IDX'],
    ['total_amount', 'DECIMAL(18,2)', '订单金额', ''],
    ['remark', 'VARCHAR(500)', '备注', ''],
    ['create_time', 'DATETIME', '创建时间', 'IDX'],
    ['create_by', 'VARCHAR(32)', '创建人', '']
]
for i, row_data in enumerate(order_fields, 1):
    for j, val in enumerate(row_data):
        order_table.rows[i].cells[j].text = val

# 卡口记录表
add_paragraph(doc, '表4：t_gate_control_record（卡口控制记录表）', bold=True)
gate_table = doc.add_table(rows=12, cols=4)
gate_table.style = 'Table Grid'
for i, h in enumerate(headers):
    cell = gate_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

gate_fields = [
    ['record_id', 'VARCHAR(32)', '记录ID', 'PK'],
    ['plate_no', 'VARCHAR(20)', '车牌号', 'IDX'],
    ['entry_time', 'DATETIME', '入场时间', 'IDX'],
    ['exit_time', 'DATETIME', '出场时间', ''],
    ['weight_in', 'DECIMAL(10,3)', '入场重量', ''],
    ['weight_out', 'DECIMAL(10,3)', '出场重量', ''],
    ['gate_status', 'TINYINT', '卡口状态', 'IDX'],
    ['customs_status', 'TINYINT', '海关状态', 'IDX'],
    ['security_status', 'TINYINT', '安检状态', 'IDX'],
    ['device_id', 'VARCHAR(32)', '设备ID', ''],
    ['create_time', 'DATETIME', '创建时间', '']
]
for i, row_data in enumerate(gate_fields, 1):
    for j, val in enumerate(row_data):
        gate_table.rows[i].cells[j].text = val

# 查验任务表
add_paragraph(doc, '表5：t_inspection_task（查验任务表）', bold=True)
inspect_table = doc.add_table(rows=12, cols=4)
inspect_table.style = 'Table Grid'
for i, h in enumerate(headers):
    cell = inspect_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

inspect_fields = [
    ['task_id', 'VARCHAR(32)', '任务ID', 'PK'],
    ['task_no', 'VARCHAR(30)', '任务编号', 'UK'],
    ['order_id', 'VARCHAR(32)', '订单ID', 'FK'],
    ['waybill_no', 'VARCHAR(30)', '运单号', 'IDX'],
    ['risk_level', 'TINYINT', '风险等级', 'IDX'],
    ['inspector_id', 'VARCHAR(32)', '查验人员ID', ''],
    ['task_status', 'TINYINT', '任务状态', 'IDX'],
    ['inspection_result', 'TINYINT', '查验结果', ''],
    ['exception_flag', 'TINYINT', '异常标记', ''],
    ['create_time', 'DATETIME', '创建时间', 'IDX'],
    ['complete_time', 'DATETIME', '完成时间', '']
]
for i, row_data in enumerate(inspect_fields, 1):
    for j, val in enumerate(row_data):
        inspect_table.rows[i].cells[j].text = val

add_paragraph(doc, '（更多数据表详见数据库设计文档，包括：运单表、仓储库存表、计费账单表、用户权限表等共30+张表）')

doc.add_page_break()

# ========== 5 中间件设计 ==========
print("生成第5章 中间件设计...")
add_heading(doc, '5 中间件设计', 1)

add_heading(doc, '5.1 消息队列（RabbitMQ）', 2)
add_paragraph(doc, '选用理由：')
add_list_item(doc, '支持消息持久化，确保消息不丢失')
add_list_item(doc, '支持镜像队列，实现高可用')
add_list_item(doc, '支持多种消息模式（直连、主题、广播）')
add_list_item(doc, '管理界面友好，易于监控和维护')
add_paragraph(doc, '应用场景：')
add_list_item(doc, '异步处理：订单创建后异步发送通知、生成账单')
add_list_item(doc, '系统解耦：海关数据申报、安检数据同步等异步处理')
add_list_item(doc, '流量削峰：高峰期订单请求削峰填谷')
add_paragraph(doc, '部署架构：3节点集群，磁盘持久化，镜像队列策略')

add_heading(doc, '5.2 缓存（Redis）', 2)
add_paragraph(doc, '选用理由：')
add_list_item(doc, '高性能读写，QPS可达10万+')
add_list_item(doc, '支持集群模式，可水平扩展')
add_list_item(doc, '数据类型丰富（String、Hash、List、Set、SortedSet）')
add_list_item(doc, '支持分布式锁、Lua脚本')
add_paragraph(doc, '应用场景：')
add_list_item(doc, '热点数据缓存：航班信息、货物状态等')
add_list_item(doc, '分布式锁：库存扣减、订单防重')
add_list_item(doc, '会话管理：用户登录会话')
add_list_item(doc, '限流计数：接口访问频次限制')
add_paragraph(doc, '部署架构：6节点集群（3主3从），启用AOF持久化')

add_heading(doc, '5.3 搜索引擎（Elasticsearch）', 2)
add_paragraph(doc, '选用理由：')
add_list_item(doc, '分布式架构，近实时搜索')
add_list_item(doc, '支持全文检索、聚合分析')
add_list_item(doc, '与ELK栈无缝集成，便于日志分析')
add_paragraph(doc, '应用场景：')
add_list_item(doc, '货物信息全文检索')
add_list_item(doc, '运单轨迹查询')
add_list_item(doc, '日志数据存储与分析')
add_paragraph(doc, '部署架构：3节点集群，分片数5，副本数1')

add_heading(doc, '5.4 对象存储（MinIO）', 2)
add_paragraph(doc, '选用理由：')
add_list_item(doc, '兼容Amazon S3 API')
add_list_item(doc, '高性能、高可用')
add_list_item(doc, '支持分布式部署')
add_paragraph(doc, '应用场景：')
add_list_item(doc, '文件存储：PDF单据、Excel报表')
add_list_item(doc, '图片存储：货物照片、查验留痕照片')
add_list_item(doc, '视频存储：监控录像')
add_paragraph(doc, '部署架构：4节点分布式集群，纠删码保护')

doc.add_page_break()

# ========== 6 人机界面设计 ==========
print("生成第6章 人机界面设计...")
add_heading(doc, '6 人机界面设计', 1)

add_heading(doc, '6.1 界面风格', 2)
add_paragraph(doc, '系统采用现代化扁平化设计风格，遵循以下设计原则：')
add_list_item(doc, '主色调：科技蓝（#1890FF），用于按钮、链接、重点标识')
add_list_item(doc, '辅助色：成功绿（#52C41A）、警告橙（#FAAD14）、错误红（#F5222D）')
add_list_item(doc, '背景色：浅灰（#F0F2F5）作为页面背景，白色作为卡片背景')
add_list_item(doc, '字体：中文使用思源黑体，英文使用Arial，正文14px，标题16-24px')
add_list_item(doc, '图标：使用Ant Design图标库，保持一致性')

add_heading(doc, '6.2 布局设计', 2)
add_paragraph(doc, 'Web管理端布局：')
add_list_item(doc, '采用左侧固定导航+右侧内容区的经典布局')
add_list_item(doc, '顶部Header包含系统Logo、消息通知、用户信息')
add_list_item(doc, '左侧Sidebar为多级菜单导航，支持收起/展开')
add_list_item(doc, '内容区采用卡片式布局，信息分组展示')
add_list_item(doc, '表格区域支持分页、筛选、排序、导出')
add_paragraph(doc, '响应式设计：')
add_list_item(doc, '支持1920x1080及以上分辨率')
add_list_item(doc, '自适应布局，适配不同屏幕尺寸')

add_heading(doc, '6.3 交互设计', 2)
add_paragraph(doc, '操作反馈：')
add_list_item(doc, '表单校验：即时校验，错误提示明确')
add_list_item(doc, '操作成功：顶部消息提示，自动消失')
add_list_item(doc, '操作失败：错误提示，提供解决建议')
add_list_item(doc, '加载状态：按钮loading、页面骨架屏')
add_paragraph(doc, '防错设计：')
add_list_item(doc, '关键操作二次确认（删除、审核）')
add_list_item(doc, '表单自动保存草稿')
add_list_item(doc, '操作撤销功能（部分场景）')

add_heading(doc, '6.4 移动端适配', 2)
add_paragraph(doc, '司机APP：')
add_list_item(doc, '简洁高效，核心功能一键触达')
add_list_item(doc, '地图集成，支持导航')
add_list_item(doc, '扫码功能，快速录入')
add_paragraph(doc, '作业PDA：')
add_list_item(doc, '大按钮设计，便于操作')
add_list_item(doc, '支持扫码枪、RFID读取')
add_list_item(doc, '离线模式，网络恢复后同步')
add_paragraph(doc, '小程序：')
add_list_item(doc, '轻量化设计，快速加载')
add_list_item(doc, '消息推送，及时通知')

doc.add_page_break()

# ========== 7 系统安全设计 ==========
print("生成第7章 系统安全设计...")
add_heading(doc, '7 系统安全设计', 1)

add_heading(doc, '7.1 身份认证', 2)
add_paragraph(doc, '认证方式：')
add_list_item(doc, '用户名/密码：密码强度要求（8位以上，含大小写字母、数字、特殊字符）')
add_list_item(doc, '短信验证码：手机号+短信验证码登录，验证码5分钟有效')
add_list_item(doc, '数字证书：与海关对接使用海关数字证书认证')
add_list_item(doc, '单点登录（SSO）：多个子系统统一认证')
add_paragraph(doc, 'Token机制：')
add_list_item(doc, '采用JWT Token，包含用户ID、角色、权限信息')
add_list_item(doc, 'Access Token有效期2小时，Refresh Token有效期7天')
add_list_item(doc, 'Token黑名单机制，支持强制登出')

add_heading(doc, '7.2 访问控制', 2)
add_paragraph(doc, 'RBAC权限模型：')
add_list_item(doc, '角色管理：系统管理员、业务管理员、操作员、查看员等')
add_list_item(doc, '权限管理：菜单权限、按钮权限、数据权限、字段权限')
add_list_item(doc, '数据范围：全部数据、本部门数据、本人数据')
add_paragraph(doc, '接口鉴权：')
add_list_item(doc, '所有接口需携带有效Token')
add_list_item(doc, '接口权限校验，无权限返回403')

add_heading(doc, '7.3 数据加密', 2)
add_paragraph(doc, '传输加密：')
add_list_item(doc, '全站HTTPS，TLS 1.2+')
add_list_item(doc, '与海关系统对接使用国密SM2/SM4加密')
add_paragraph(doc, '存储加密：')
add_list_item(doc, '敏感字段加密存储（密码、身份证号、手机号）')
add_list_item(doc, '使用AES-256加密算法')
add_list_item(doc, '密钥托管在KMS（密钥管理系统）')

add_heading(doc, '7.4 安全审计', 2)
add_list_item(doc, '登录审计：记录登录时间、IP地址、设备信息、登录结果')
add_list_item(doc, '操作审计：记录增删改操作，包含操作前数据快照')
add_list_item(doc, '数据变更审计：敏感数据变更记录，支持数据追溯')
add_list_item(doc, '审计日志保留180天，定期归档')

add_heading(doc, '7.5 漏洞防护', 2)
add_list_item(doc, 'SQL注入防护：使用MyBatis参数化查询，禁止字符串拼接SQL')
add_list_item(doc, 'XSS攻击防护：输入过滤、输出转义、CSP策略')
add_list_item(doc, 'CSRF攻击防护：Token校验、SameSite Cookie')
add_list_item(doc, '文件上传安全：类型白名单、大小限制、病毒扫描、重命名存储')

doc.add_page_break()

# ========== 8 业务连续性设计 ==========
print("生成第8章 业务连续性设计...")
add_heading(doc, '8 业务连续性设计', 1)

add_heading(doc, '8.1 高可用设计', 2)
add_paragraph(doc, '应用层高可用：')
add_list_item(doc, '无状态设计，应用服务可多节点部署')
add_list_item(doc, 'Kubernetes自动故障转移，Pod异常自动重启')
add_list_item(doc, '负载均衡：Nginx/Traefik，支持健康检查')
add_paragraph(doc, '数据库高可用：')
add_list_item(doc, 'MySQL主从复制，主库故障自动切换')
add_list_item(doc, 'Redis集群，节点故障自动故障转移')
add_list_item(doc, 'RabbitMQ镜像队列，节点故障不影响消息服务')

add_heading(doc, '8.2 容灾备份', 2)
add_paragraph(doc, '数据备份策略：')
add_list_item(doc, '数据库：每日全量备份+每小时增量备份，保留30天')
add_list_item(doc, '对象存储：多副本存储，跨可用区冗余')
add_list_item(doc, '配置文件：Git版本控制，变更历史可追溯')
add_list_item(doc, '备份数据异地存储，距离大于100公里')
add_paragraph(doc, '灾难恢复：')
add_list_item(doc, 'RTO（恢复时间目标）< 30分钟')
add_list_item(doc, 'RPO（恢复点目标）< 15分钟')

add_heading(doc, '8.3 故障恢复', 2)
add_list_item(doc, '自动故障检测：Prometheus监控告警，异常自动通知')
add_list_item(doc, '自动故障转移：数据库主从切换、服务实例切换')
add_list_item(doc, '手动切换预案：关键故障人工确认后执行切换')
add_list_item(doc, '定期演练：每季度进行一次故障演练')

doc.add_page_break()

print("第4-8章已生成，保存文档...")

# 保存文档
doc.save('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')
print("文档更新完成！")
