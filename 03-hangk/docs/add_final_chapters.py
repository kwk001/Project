#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
补充技术方案第9-14章详细内容
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading

def add_paragraph(doc, text, bold=False, indent=True):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    return para

def add_list_item(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    return para

# 打开已有文档
doc = Document('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')

# ========== 9 系统接口设计 ==========
print("生成第9章 系统接口设计...")
add_heading(doc, '9 系统接口设计', 1)

add_heading(doc, '9.1 接口总表', 2)
interface_table = doc.add_table(rows=7, cols=6)
interface_table.style = 'Table Grid'
headers = ['序号', '生产系统', '消费系统', '接口方向', '协议类型', '备注']
for i, h in enumerate(headers):
    cell = interface_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

interface_data = [
    ['1', '海关金关二期', '本平台', '双向', 'HTTPS/JSON', '海关业务数据'],
    ['2', '机场安检系统', '本平台', '双向', 'WebService/XML', '安检数据'],
    ['3', '场站系统', '本平台', '接收', 'REST/JSON', '航班数据'],
    ['4', '称重设备', '本平台', '接收', 'TCP/Modbus', '称重数据'],
    ['5', '车牌识别设备', '本平台', '接收', 'SDK/API', '识别结果'],
    ['6', '分拣线PLC', '本平台', '双向', 'OPC UA', '控制指令']
]
for i, row_data in enumerate(interface_data, 1):
    for j, val in enumerate(row_data):
        interface_table.rows[i].cells[j].text = val

add_paragraph(doc, '说明：本表是根据招标文件和用户需求调研成果文件梳理的本系统现阶段的接口汇总，后期根据项目进展和对接情况进行调整和优化，详细接口内容详见各业务系统签订的接口协议。')

add_heading(doc, '9.2 海关金关二期系统接口', 2)
add_paragraph(doc, '接口概述：', bold=True)
add_list_item(doc, '接口方向：双向')
add_list_item(doc, '通信协议：HTTPS + JSON')
add_list_item(doc, '认证方式：海关数字证书 + 报文签名验证')
add_list_item(doc, '字符编码：UTF-8')

add_paragraph(doc, '接口功能：')
add_list_item(doc, '货物预报信息报送：将货物基础信息、收发货人信息、航班信息报送至海关')
add_list_item(doc, '运抵信息报送：货物入站后向海关报送运抵信息')
add_list_item(doc, '查验结果报送：查验完成后向海关报送查验结果')
add_list_item(doc, '放行指令接收：接收海关系统的放行/查验/扣留指令')
add_list_item(doc, '风险预警接收：接收海关系统的风险预警信息')

add_paragraph(doc, '请求报文示例：', bold=True)
add_paragraph(doc, '''{
  "messageHeader": {
    "messageId": "MSG202404080001",
    "messageType": "DECLARE",
    "sendTime": "2024-04-08T10:30:00+08:00",
    "sender": "AIRPORT_PLAT",
    "receiver": "CUSTOMS_GJ2",
    "sign": "xxxxxxxxxxxxxxxx"
  },
  "messageBody": {
    "declareNo": "DEC202404080001",
    "waybillNo": "999-12345675",
    "cargoName": "电子产品",
    "cargoWeight": 150.5,
    "cargoQuantity": 10,
    "cargoValue": 50000.00,
    "currency": "CNY",
    "senderInfo": {
      "name": "发货人名称",
      "code": "SENDER001"
    },
    "receiverInfo": {
      "name": "收货人名称",
      "country": "US"
    }
  }
}''', indent=False)

add_paragraph(doc, '响应报文示例：', bold=True)
add_paragraph(doc, '''{
  "code": "0000",
  "message": "接收成功",
  "data": {
    "declareNo": "DEC202404080001",
    "customsNo": "C2024040800001",
    "status": "ACCEPTED",
    "receiveTime": "2024-04-08T10:30:02+08:00"
  }
}''', indent=False)

add_heading(doc, '9.3 机场安检系统接口', 2)
add_paragraph(doc, '接口概述：', bold=True)
add_list_item(doc, '接口方向：双向')
add_list_item(doc, '通信协议：WebService + XML')
add_list_item(doc, '认证方式：IP白名单 + Token认证')

add_paragraph(doc, '接口功能：')
add_list_item(doc, '货物安检申报：货物入站前向安检系统申报')
add_list_item(doc, '安检结果接收：接收安检系统的安检结果（通过/不通过）')
add_list_item(doc, '安检图像获取：获取安检设备的扫描图像')
add_list_item(doc, '安检状态同步：实时同步安检状态')

add_heading(doc, '9.4 场站系统接口', 2)
add_paragraph(doc, '接口概述：', bold=True)
add_list_item(doc, '接口方向：本平台接收')
add_list_item(doc, '通信协议：REST API + JSON')
add_list_item(doc, '认证方式：OAuth 2.0')

add_paragraph(doc, '接口功能：')
add_list_item(doc, '航班动态信息接收：接收航班计划变更、起飞/到达状态')
add_list_item(doc, '货物装机信息接收：接收货物装机确认信息')
add_list_item(doc, '货物离港信息接收：接收航班离港信息')
add_list_item(doc, '航班状态变更通知：实时接收航班状态变更')

add_heading(doc, '9.5 设备接口', 2)
add_paragraph(doc, '9.5.1 智能称重设备接口', bold=True)
add_list_item(doc, '通信协议：TCP/Modbus RTU')
add_list_item(doc, '数据格式：称重数据（重量、时间、设备号）')
add_list_item(doc, '采集频率：实时采集，车辆上磅稳定后触发')

add_paragraph(doc, '9.5.2 车牌识别设备接口', bold=True)
add_list_item(doc, '通信协议：设备SDK/API')
add_list_item(doc, '数据格式：车牌号、识别时间、置信度、车辆图片')
add_list_item(doc, '识别要求：准确率>99%，响应时间<1秒')

add_paragraph(doc, '9.5.3 分拣线PLC接口', bold=True)
add_list_item(doc, '通信协议：OPC UA')
add_list_item(doc, '数据交换：控制指令下发、状态读取、故障报警')
add_list_item(doc, '实时性要求：控制指令响应时间<100ms')

doc.add_page_break()

# ========== 10 工程界面设计 ==========
print("生成第10章 工程界面设计...")
add_heading(doc, '10 工程界面设计', 1)

add_paragraph(doc, '【与海关金关二期系统的工程界面】', bold=True)
add_paragraph(doc, '物理边界：')
add_list_item(doc, '海关专网与本平台应用服务器之间通过网闸进行物理隔离')
add_list_item(doc, '网闸部署在海关专网边界的数据交换区')
add_list_item(doc, '数据交换采用"摆渡"方式，内网与外网无直接连接')

add_paragraph(doc, '责任划分：')
add_list_item(doc, '本平台责任：' )
add_list_item(doc, '  - 负责向海关系统报送货物预报、运抵、查验等数据', 1)
add_list_item(doc, '  - 负责接收海关系统的放行/查验/风险预警指令', 1)
add_list_item(doc, '  - 负责数据的格式转换和协议适配', 1)
add_list_item(doc, '海关系统责任：')
add_list_item(doc, '  - 负责接收并处理本平台报送的数据', 1)
add_list_item(doc, '  - 负责下发监管指令', 1)
add_list_item(doc, '  - 负责指令的合法性和准确性', 1)

add_paragraph(doc, '接口位置：')
add_list_item(doc, '部署在海关专网边界的数据交换区')
add_list_item(doc, '本平台侧通过网闸外网接口接入')
add_list_item(doc, '海关系统侧通过网闸内网接口接入')

add_paragraph(doc, '【与机场安检系统的工程界面】', bold=True)
add_paragraph(doc, '物理边界：')
add_list_item(doc, '安检系统网络与货站网络通过防火墙隔离')
add_list_item(doc, '接口服务器部署在DMZ区')

add_paragraph(doc, '责任划分：')
add_list_item(doc, '本平台负责货物信息报送和安检结果接收')
add_list_item(doc, '安检系统负责安检执行和结果反馈')

doc.add_page_break()

# ========== 11 软硬件资源需求分析 ==========
print("生成第11章 软硬件资源需求分析...")
add_heading(doc, '11 软硬件资源需求分析', 1)

add_heading(doc, '11.1 云资源需求', 2)
resource_table = doc.add_table(rows=7, cols=5)
resource_table.style = 'Table Grid'
headers = ['资源类型', '配置', '数量', '用途', '备注']
for i, h in enumerate(headers):
    cell = resource_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

resource_data = [
    ['应用服务器', '8C16G', '4台', '业务应用', '负载均衡部署'],
    ['数据库服务器', '16C32G', '2台', 'MySQL主从', 'SSD存储500G'],
    ['缓存服务器', '4C8G', '3台', 'Redis集群', '主从+哨兵'],
    ['ES服务器', '8C16G', '3台', '搜索引擎', '3节点集群'],
    ['对象存储', '标准存储', '1个', '文件存储', '容量1TB起步'],
    ['消息队列', '4C8G', '2台', 'RabbitMQ', '镜像队列']
]
for i, row_data in enumerate(resource_data, 1):
    for j, val in enumerate(row_data):
        resource_table.rows[i].cells[j].text = val

add_heading(doc, '11.2 网络需求', 2)
add_paragraph(doc, '带宽需求：')
add_list_item(doc, '互联网带宽：100Mbps，用于公众用户移动端访问')
add_list_item(doc, '海关专线带宽：50Mbps，用于与海关系统数据交换')
add_list_item(doc, '内网带宽：1Gbps，用于系统内部服务间通信')
add_list_item(doc, '公网IP：4个，用于负载均衡和网关')

add_paragraph(doc, '网络安全：')
add_list_item(doc, 'WAF：Web应用防火墙，防护SQL注入、XSS等攻击')
add_list_item(doc, 'DDoS防护：10Gbps防护能力')
add_list_item(doc, '防火墙：南北向和东西向流量隔离')
add_list_item(doc, 'VPN：运维人员远程接入')

doc.add_page_break()

# ========== 12 集成测试设计 ==========
print("生成第12章 集成测试设计...")
add_heading(doc, '12 集成测试设计', 1)

add_paragraph(doc, '测试策略：')
add_list_item(doc, '测试范围：四大子系统间的接口集成、与外部系统的对接集成、端到端业务流程')
add_list_item(doc, '测试方法：自下而上集成，先单元测试，后接口测试，再端到端测试')
add_list_item(doc, '测试环境：独立测试环境，数据与生产隔离')
add_list_item(doc, '自动化程度：接口测试自动化率>80%，UI测试自动化率>60%')

add_paragraph(doc, '测试内容：')
add_paragraph(doc, '1. 系统集成测试', bold=True)
add_list_item(doc, '航空货运与海关系统集成：订单数据同步、查验结果同步')
add_list_item(doc, '海关系统与卡口设备集成：车牌识别、称重数据采集、道闸控制')
add_list_item(doc, '园区系统与安防设备集成：摄像头、门禁、消防传感器')

add_paragraph(doc, '2. 接口集成测试', bold=True)
add_list_item(doc, '海关金关二期接口：数据报送准确性、指令接收及时性')
add_list_item(doc, '安检系统接口：数据同步、状态反馈')
add_list_item(doc, '场站系统接口：航班数据接收、货物状态同步')

add_paragraph(doc, '3. 业务流程测试', bold=True)
add_list_item(doc, '货物入站流程：从预约到入库的完整流程')
add_list_item(doc, '货物出站流程：从出库申请到上机确认的完整流程')
add_list_item(doc, '海关查验流程：从AI筛选到结果同步的完整流程')
add_list_item(doc, '跨境电商流程：从订单申报到结关反馈的完整流程')

add_paragraph(doc, '通过标准：')
add_list_item(doc, '功能覆盖率100%，所有功能点均有测试用例覆盖')
add_list_item(doc, '致命缺陷0个，严重缺陷<3个，一般缺陷<10个')
add_list_item(doc, '接口成功率>99.9%，平均响应时间<500ms')

doc.add_page_break()

# ========== 13 演练方案设计 ==========
print("生成第13章 演练方案设计...")
add_heading(doc, '13 演练方案设计', 1)

add_heading(doc, '13.1 演练场景', 2)

add_paragraph(doc, '场景1：系统压力演练', bold=True)
add_list_item(doc, '演练目标：验证系统在高并发场景下的承载能力')
add_list_item(doc, '演练内容：模拟双十一高峰期业务量（日常10倍），持续2小时')
add_list_item(doc, '评估指标：系统响应时间、错误率、资源利用率')
add_list_item(doc, '通过标准：响应时间<1秒，错误率<0.1%，CPU利用率<80%')

add_paragraph(doc, '场景2：故障切换演练', bold=True)
add_list_item(doc, '演练目标：验证系统故障自动切换能力')
add_list_item(doc, '演练内容：模拟主数据库故障、Redis节点故障、应用Pod故障')
add_list_item(doc, '评估指标：故障检测时间、切换时间、业务影响时间')
add_list_item(doc, '通过标准：RTO<5分钟，数据零丢失')

add_paragraph(doc, '场景3：安全应急演练', bold=True)
add_list_item(doc, '演练目标：验证系统安全防护能力')
add_list_item(doc, '演练内容：模拟DDoS攻击、SQL注入尝试、暴力破解')
add_list_item(doc, '评估指标：攻击检测率、防护成功率、响应时间')
add_list_item(doc, '通过标准：攻击全部拦截，系统正常运行')

add_paragraph(doc, '场景4：业务应急演练', bold=True)
add_list_item(doc, '演练目标：验证海关系统中断时的本地应急处理能力')
add_list_item(doc, '演练内容：模拟海关系统网络中断，验证本地模式运行')
add_list_item(doc, '评估指标：本地模式切换时间、数据缓存能力、恢复后同步能力')
add_list_item(doc, '通过标准：切换时间<1分钟，数据无丢失，恢复后自动同步')

add_heading(doc, '13.2 演练计划', 2)
add_paragraph(doc, '演练频次：每季度至少一次综合演练')
add_paragraph(doc, '参与人员：运维团队、开发团队、业务人员、安全人员')
add_paragraph(doc, '演练报告：每次演练后输出演练报告，记录问题并跟踪整改')

doc.add_page_break()

# ========== 14 试运行方案设计 ==========
print("生成第14章 试运行方案设计...")
add_heading(doc, '14 试运行方案设计', 1)

add_heading(doc, '14.1 试运行范围', 2)
add_paragraph(doc, '第一阶段（1个月）：')
add_list_item(doc, '选取1个货站作为试点')
add_list_item(doc, '上线航空货运基础功能和海关基础功能')
add_list_item(doc, '验证核心业务流程')

add_paragraph(doc, '第二阶段（1个月）：')
add_list_item(doc, '扩展至全园区所有货站')
add_list_item(doc, '上线全部功能模块')
add_list_item(doc, '全量业务压力验证')

add_paragraph(doc, '第三阶段（1个月）：')
add_list_item(doc, '系统优化调整')
add_list_item(doc, '性能调优')
add_list_item(doc, '用户培训')

add_heading(doc, '14.2 切换策略', 2)
add_paragraph(doc, '并行运行期（1个月）：')
add_list_item(doc, '新旧系统并行运行，业务数据双写')
add_list_item(doc, '新系统处理主要业务，旧系统作为备份')
add_list_item(doc, '每日核对数据一致性')

add_paragraph(doc, '切流策略：')
add_list_item(doc, '按业务模块逐步切流，从非核心模块开始')
add_list_item(doc, '切流前进行数据完整性校验')
add_list_item(doc, '切流后实时监控，发现问题立即回切')

add_heading(doc, '14.3 回退方案', 2)
add_paragraph(doc, '回退触发条件：')
add_list_item(doc, '系统出现致命缺陷，影响业务正常开展')
add_list_item(doc, '数据不一致或数据丢失')
add_list_item(doc, '性能严重下降，无法满足业务需求')

add_paragraph(doc, '回退步骤：')
add_list_item(doc, '步骤1：立即停止新系统业务受理')
add_list_item(doc, '步骤2：切换至旧系统继续业务')
add_list_item(doc, '步骤3：核对新旧系统数据差异')
add_list_item(doc, '步骤4：补齐差异数据')
add_list_item(doc, '步骤5：问题修复后重新切流')

add_paragraph(doc, '回退保障：')
add_list_item(doc, '保留旧系统运行环境3个月')
add_list_item(doc, '制定详细的回退操作手册')
add_list_item(doc, '进行回退演练，确保可执行')

print("第9-14章已生成，保存文档...")

# 保存文档
doc.save('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')
print("完整文档生成完成！")
print("文件路径: 国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx")
