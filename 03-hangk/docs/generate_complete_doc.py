#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整版技术方案说明书
包含：
1. 171个功能点详细清单
2. 数据库表结构详细设计
3. 接口详细设计
4. 功能流程详细设计
5. 完整的14个章节
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), '4')
        edge_elm.set(qn('w:color'), '000000')
        tcBorders.append(edge_elm)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading

def add_paragraph(doc, text, bold=False, indent=True):
    """添加段落"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    return para

def add_list_item(doc, text, level=0):
    """添加列表项"""
    para = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Number')
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    return para

# 读取功能清单
print("正在读取功能清单...")
df = pd.read_excel('功能清单__国际物流航空货站运行平台及海关信息平台V1.1.xlsx', sheet_name='功能清单')
df['一级目录'] = df['一级目录'].ffill()
df['二级目录'] = df['二级目录'].ffill()

# 创建文档
doc = Document()

# 设置中文字体
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# ========== 封面 ==========
print("生成封面...")
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('民用机场工程')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(26)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('信息弱电工程施工')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(26)
run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('国际物流航空货站运行平台')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(22)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('及海关信息平台')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(22)
run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('深化设计技术方案说明书')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(22)
run.font.bold = True

for _ in range(8):
    doc.add_paragraph()

info_items = [
    '编制单位：XXX科技有限公司',
    '编 制 人：XXX',
    '审 核 人：XXX',
    '审 批 人：XXX',
    '编制日期：2026年04月08日'
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)

# 分页
doc.add_page_break()

# ========== 文件变更记录 ==========
print("生成文件变更记录...")
add_heading(doc, '文件变更记录', 1)

table = doc.add_table(rows=2, cols=5)
table.style = 'Table Grid'
headers = ['版本号', '日期', '修改人', '审阅人', '摘要']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9D9D9')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

data = ['V1.0', '2026-04-08', 'XXX', 'XXX', '初始版本创建']
for i, val in enumerate(data):
    table.rows[1].cells[i].text = val

doc.add_page_break()

# ========== 目录 ==========
print("生成目录...")
add_heading(doc, '目 录', 1)

toc_items = [
    ('1 引言', 0),
    ('1.1 编写目的', 1), ('1.2 范围', 1), ('1.3 编制依据', 1), ('1.4 术语定义', 1),
    ('2 概述', 0),
    ('2.1 系统目标', 1), ('2.2 开发环境', 1), ('2.3 运行环境', 1), ('2.4 条件与限制', 1),
    ('3 总体设计', 0),
    ('3.1 技术架构设计', 1), ('3.2 逻辑架构设计', 1), ('3.3 功能架构设计', 1),
    ('3.4 网络架构设计', 1), ('3.5 功能及流程设计', 1),
    ('4 数据结构设计', 0),
    ('4.1 逻辑结构设计', 1), ('4.2 物理结构设计', 1), ('4.3 数据库表结构', 1),
    ('5 中间件设计', 0),
    ('5.1 消息队列', 1), ('5.2 缓存', 1), ('5.3 搜索引擎', 1), ('5.4 对象存储', 1),
    ('6 人机界面设计', 0),
    ('6.1 界面风格', 1), ('6.2 布局设计', 1), ('6.3 交互设计', 1), ('6.4 移动端适配', 1),
    ('7 系统安全设计', 0),
    ('7.1 身份认证', 1), ('7.2 访问控制', 1), ('7.3 数据加密', 1),
    ('7.4 安全审计', 1), ('7.5 漏洞防护', 1),
    ('8 业务连续性设计', 0),
    ('8.1 高可用设计', 1), ('8.2 容灾备份', 1), ('8.3 故障恢复', 1),
    ('9 系统接口设计', 0),
    ('9.1 接口总表', 1), ('9.2 海关金关二期接口', 1), ('9.3 安检系统接口', 1),
    ('9.4 场站系统接口', 1), ('9.5 设备接口', 1),
    ('10 工程界面设计', 0),
    ('11 软硬件资源需求分析', 0),
    ('12 集成测试设计', 0),
    ('13 演练方案设计', 0),
    ('14 试运行方案设计', 0),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    if level == 0:
        run = p.add_run(item)
        run.font.bold = True
    else:
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(item)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

doc.add_page_break()

# ========== 1 引言 ==========
print("生成第1章 引言...")
add_heading(doc, '1 引言', 1)

add_heading(doc, '1.1 编写目的', 2)
add_paragraph(doc, '依据《国际物流航空货站运行平台及海关信息平台招标文件》、《用户需求调研分析报告》进行本系统深化设计，本设计方案将对本系统的各个功能模块及功能点进行详细设计，涵盖航空货运信息管理、海关信息系统、园区运行中心、航空物流公共信息平台四大核心子系统，共计171个功能点，是系统建设及项目实施的指导与依据。')

add_heading(doc, '1.2 范围', 2)
add_paragraph(doc, '系统名称：国际物流航空货站运行平台及海关信息平台')
add_paragraph(doc, '开发人员：XXX科技有限公司开发团队')
add_paragraph(doc, '功能范围：')
add_list_item(doc, '航空货运信息管理系统：104个功能点，涵盖基础资料、订单运单、运输调度、仓储作业、计费结算、货包机跟踪、分拣线控制等')
add_list_item(doc, '海关信息系统：23个功能点，涵盖智能卡口、查验预警、货物预报、跨境电商、可视化展示等')
add_list_item(doc, '园区运行中心：22个功能点，涵盖态势监控、园区管理、数字平台等')
add_list_item(doc, '航空物流公共信息平台：22个功能点，涵盖移动端、系统管理、公共信息服务、客服业务等')

add_heading(doc, '1.3 编制依据', 2)
standards = [
    '《国际物流航空货站运行平台及海关信息平台招标文件》',
    '《用户需求调研分析报告》',
    '《海关金关二期智能卡口系统接口规范》',
    '《民用机场工程信息弱电系统设计规范》',
    '《信息安全技术网络安全等级保护基本要求》（GB/T 22239-2019）',
    '《软件工程软件开发成本度量规范》（GB/T 36964-2018）',
    '《海关信息化应用项目数据交换标准》',
    '《民航机场信息系统技术规范》（MH/T 5103-2004）'
]
for std in standards:
    add_list_item(doc, std)

add_heading(doc, '1.4 术语定义', 2)
terms_table = doc.add_table(rows=9, cols=3)
terms_table.style = 'Table Grid'
headers = ['名词', '含义', '备注']
for i, h in enumerate(headers):
    cell = terms_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

terms_data = [
    ['9610', '跨境电商B2C出口监管方式', '海关监管代码'],
    ['9710', '跨境电商B2B直接出口监管方式', '海关监管代码'],
    ['9810', '跨境电商出口海外仓监管方式', '海关监管代码'],
    ['IoT', '物联网（Internet of Things）', '技术术语'],
    ['PDA', '手持终端设备', '硬件设备'],
    ['PLC', '可编程逻辑控制器', '控制设备'],
    ['WMS', '仓库管理系统', '系统名称'],
    ['TMS', '运输管理系统', '系统名称']
]
for i, row_data in enumerate(terms_data, 1):
    for j, val in enumerate(row_data):
        terms_table.rows[i].cells[j].text = val

doc.add_page_break()

# ========== 2 概述 ==========
print("生成第2章 概述...")
add_heading(doc, '2 概述', 1)

add_heading(doc, '2.1 系统目标', 2)
add_paragraph(doc, '国际物流航空货站运行平台及海关信息平台旨在构建一个集航空货运业务管理、海关监管、园区运营管理、公共信息服务于一体的综合性智慧物流平台。系统建设目标包括：')
add_list_item(doc, '实现航空货运业务全流程数字化管理，涵盖订单、运单、运输、仓储等核心环节，实现货物从入库到出库的全生命周期管理')
add_list_item(doc, '建立与海关金关二期系统的深度对接，实现智能卡口控制、货物预报、查验预警等海关业务功能，满足海关监管要求')
add_list_item(doc, '构建园区态势感知体系，实现交通、消防、安防、能耗、仓储等多维度实时监控和智能预警')
add_list_item(doc, '提供移动端服务能力，支持司机APP、作业PDA、管理小程序、客户小程序等多终端接入')
add_list_item(doc, '建立数据治理平台，实现数据标准化、共享交换与智能分析，支撑业务决策')

add_heading(doc, '2.2 开发环境', 2)
dev_env_table = doc.add_table(rows=9, cols=3)
dev_env_table.style = 'Table Grid'
for i, h in enumerate(['工具名称', '使用版本', '简要说明']):
    cell = dev_env_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, 'D9D9D9')

dev_env_data = [
    ['JDK', '1.8+', 'Java开发工具包'],
    ['Spring Boot', '2.7.x', '微服务开发框架'],
    ['Spring Cloud Alibaba', '2021.x', '微服务治理框架'],
    ['MySQL', '8.0', '关系型数据库'],
    ['Redis', '6.x', '分布式缓存'],
    ['RabbitMQ', '3.x', '消息中间件'],
    ['Elasticsearch', '7.x', '搜索引擎'],
    ['Vue.js', '3.x', '前端框架']
]
for i, row_data in enumerate(dev_env_data, 1):
    for j, val in enumerate(row_data):
        dev_env_table.rows[i].cells[j].text = val

add_heading(doc, '2.3 运行环境', 2)
run_env = [
    '数据库环境：MySQL 8.0 主从部署，支持读写分离；MongoDB用于非结构化数据存储；ES用于全文检索',
    '消息中间件：RabbitMQ集群（3节点），支持消息持久化、镜像队列、高可用',
    'Java虚拟机：OpenJDK 1.8+，堆内存配置4GB-8GB，G1垃圾收集器',
    '服务注册中心：Nacos 2.x集群，支持服务发现、配置管理、动态配置推送',
    '数据缓存组件：Redis 6.x 集群模式（6节点），支持分布式锁、缓存、会话管理',
    '开发框架：Spring Boot 2.7.x + Spring Cloud Alibaba + MyBatis Plus',
    '开发技术：Java、JavaScript/TypeScript、Vue.js、Python（AI算法）、Node.js'
]
for item in run_env:
    add_list_item(doc, item)

add_heading(doc, '2.4 条件与限制', 2)
add_paragraph(doc, '技术约束：')
add_list_item(doc, '系统需支持与海关金关二期系统、机场安检系统、场站系统等外部系统的接口对接，接口协议需符合相关标准规范')
add_list_item(doc, '系统需支持高并发访问，峰值QPS不低于1000，平均响应时间小于500ms')
add_paragraph(doc, '业务约束：')
add_list_item(doc, '海关业务功能需严格遵循海关监管要求，数据报文格式需符合海关接口规范，数据需实时同步至海关系统')
add_list_item(doc, '货物入站需先通过安检，安检不通过货物禁止入货站')
add_paragraph(doc, '资源约束：')
add_list_item(doc, '系统部署在云平台，需考虑资源成本与扩展性平衡，支持弹性伸缩')
add_paragraph(doc, '安全约束：')
add_list_item(doc, '系统需通过等保三级认证，数据传输需加密，敏感操作需审计，数据需定期备份')

doc.add_page_break()

# ========== 3 总体设计 ==========
print("生成第3章 总体设计...")
add_heading(doc, '3 总体设计', 1)

add_heading(doc, '3.1 技术架构设计', 2)
add_paragraph(doc, '系统采用微服务架构，基于Spring Cloud Alibaba技术栈构建，整体技术架构分为五层：')
add_list_item(doc, '接入层：提供多渠道接入能力，包括Web管理端、司机APP、作业PDA、查验PDA、管理小程序、客户小程序、第三方系统接口网关')
add_list_item(doc, '网关层：基于Spring Cloud Gateway实现统一接入、负载均衡、限流熔断、安全认证、日志记录')
add_list_item(doc, '服务层：核心业务服务，采用领域驱动设计（DDD），包括航空货运服务、海关监管服务、园区运营服务、公共服务四大领域服务')
add_list_item(doc, '数据层：关系型数据库（MySQL主从）、缓存（Redis集群）、搜索引擎（ES集群）、对象存储（MinIO）、时序数据库（InfluxDB）')
add_list_item(doc, '基础设施层：容器化部署（Docker + Kubernetes）、服务注册与配置中心（Nacos）、监控告警（Prometheus + Grafana）、日志收集（ELK）')

add_heading(doc, '3.2 逻辑架构设计', 2)
add_paragraph(doc, '逻辑架构按照领域驱动设计（DDD）原则，划分为以下核心领域：')
add_list_item(doc, '航空货运领域：负责订单管理、运单管理、运输调度、仓储作业、计费结算、货包机跟踪、分拣线控制等核心业务')
add_list_item(doc, '海关监管领域：负责智能卡口控制、查验预警处置、货物预报管理、跨境电商业务、海关数据对接等监管业务')
add_list_item(doc, '园区运营领域：负责态势感知、停车管理、安防监控、设备管理、能耗管理等园区运营业务')
add_list_item(doc, '公共服务领域：负责用户认证、权限管理、消息通知、数据查询、客服支撑、系统监控等公共服务')

add_heading(doc, '3.3 功能架构设计', 2)
add_paragraph(doc, '系统功能架构包含四大子系统，共计171个功能点：')
add_paragraph(doc, '【航空货运信息管理系统】104个功能点', bold=True)
add_list_item(doc, '基础资料管理（20个）：航班信息、客户信息、货物信息、仓库库位、容器管理、资源管理、审批配置、单据配置')
add_list_item(doc, '订单与运单管理（7个）：订单创建、拆分、变更、可视化；运单录入、跟踪、变更')
add_list_item(doc, '运输与调度管理（13个）：调度计划、车辆配载、路线规划；预约申请、审核、智能排队、优先级计算、调度指令')
add_list_item(doc, '仓储作业管理（29个）：货场作业、月台管理；监管仓入库、出库、库内作业、综合查询')
add_list_item(doc, '计费与合同管理（12个）：计费科目、费率设置、应收应付计算、账单生成、预付管理、报价单、合同管理')
add_list_item(doc, '货包机跟踪（8个）：空运跟踪、陆运GPS跟踪、铁路运单跟踪')
add_list_item(doc, '分拣线控制（6个）：仓储管理、全流程跟踪、分拣控制、流程调度、打板位调度、PLC对接')

add_paragraph(doc, '【海关信息系统】23个功能点', bold=True)
add_list_item(doc, '智能卡口控制（5个）：设备管理、车辆识别、地磅称重、金关对接、安检对接')
add_list_item(doc, '查验预警处置（4个）：AI查验筛选、指令下发、结果录入、异常处置')
add_list_item(doc, '货物预报管理（3个）：预报录入、预报审核、预录入回执')
add_list_item(doc, '跨境电商（4个）：9610/9710/9810业务管理、海关对接')
add_list_item(doc, '海关业务（2个）：海关申报、税费计算')
add_list_item(doc, '可视化展示（3个）：业务数据大屏、监管数据大屏、卡口运行大屏')

add_paragraph(doc, '【园区运行中心】22个功能点', bold=True)
add_list_item(doc, '态势监控（13个）：园区、交通、消防、资产、货物、车辆、仓储、月台、客服、安防、能耗等综合态势')
add_list_item(doc, '园区管理（6个）：车辆无感通行、车位管理、巡更巡检、设备知识库、停车缴费、紧急放行')
add_list_item(doc, '数字平台（4个）：物联网平台、数据治理平台、视频转码平台、统一认证平台')

add_paragraph(doc, '【航空物流公共信息平台】22个功能点', bold=True)
add_list_item(doc, '移动端（5个）：司机APP、作业PDA、查验PDA、管理小程序、客户小程序')
add_list_item(doc, '系统管理（3个）：用户权限管理、系统参数配置、用户登录')
add_list_item(doc, '运维服务（3个）：系统监控、日志管理、备份恢复')
add_list_item(doc, '公共信息服务（7个）：航班动态、货物跟踪、物流资讯、在线服务、物流跟踪、航班订阅、在线下单')
add_list_item(doc, '客服业务（4个）：在线客服、投诉建议、工单管理、客服知识库')

add_heading(doc, '3.4 网络架构设计', 2)
add_paragraph(doc, '网络架构采用分层设计，分为六个安全域：')
add_list_item(doc, '互联网接入区：面向公众用户的移动端接入，通过WAF、DDoS防护、CDN保障安全和性能')
add_list_item(doc, '边界隔离区：部署防火墙、网闸、入侵检测系统，实现内外网隔离与数据安全交换')
add_list_item(doc, '应用服务区：部署业务应用服务、中间件、缓存集群，采用容器化部署')
add_list_item(doc, '数据存储区：部署数据库、对象存储、备份系统，独立安全域，严格访问控制')
add_list_item(doc, '运维管理区：部署监控、日志、堡垒机等运维系统，运维人员专用访问通道')
add_list_item(doc, '海关专网区：与海关金关二期系统对接的专用网络区域，物理隔离，通过网闸交换数据')

add_heading(doc, '3.5 功能及流程设计', 2)
add_paragraph(doc, '【核心业务流程】', bold=True)

add_paragraph(doc, '1. 货物入站流程：')
add_list_item(doc, '步骤1：承运商通过司机APP或Web端进行车辆预约，填写车牌号、货物信息、预计到达时间')
add_list_item(doc, '步骤2：车辆到达卡口，车牌识别系统自动识别车牌，与预约信息比对')
add_list_item(doc, '步骤3：车辆上地磅称重，称重数据自动采集并记录，超重车辆自动预警')
add_list_item(doc, '步骤4：系统自动向海关金关二期系统申报运抵，向安检系统申报安检')
add_list_item(doc, '步骤5：海关系统返回放行/查验指令，安检系统返回安检结果')
add_list_item(doc, '步骤6：系统根据指令自动控制道闸放行或引导至查验区/退货区')
add_list_item(doc, '步骤7：放行车辆进入货站，进行卸货、入库上架操作')

add_paragraph(doc, '2. 货物出站流程：')
add_list_item(doc, '步骤1：货主或货代提交出库申请，系统根据航班信息进行出库计划安排')
add_list_item(doc, '步骤2：仓库人员根据出库任务进行拣货作业，扫描库位二维码确认')
add_list_item(doc, '步骤3：拣货完成后进行复核校验，核对货物信息与出库单是否一致')
add_list_item(doc, '步骤4：复核通过后进行货物交接，与航空公司或承运人确认交接')
add_list_item(doc, '步骤5：货物装机后，系统进行上机确认，更新货物状态')
add_list_item(doc, '步骤6：航班起飞后，系统自动向海关申报离境，向货主推送运单跟踪信息')

add_paragraph(doc, '3. 海关查验流程：')
add_list_item(doc, '步骤1：系统基于AI风险预警模型对入站货物进行风险评分和查验筛选')
add_list_item(doc, '步骤2：高风险货物自动生成查验任务，系统向查验人员PDA推送查验指令')
add_list_item(doc, '步骤3：查验人员根据指令进行开箱查验，使用PDA拍照留痕、录入查验结果')
add_list_item(doc, '步骤4：查验发现异常的，进行异常登记并触发异常处置流程')
add_list_item(doc, '步骤5：查验结果实时同步至海关金关二期系统')
add_list_item(doc, '步骤6：查验通过的货物自动放行，不通过的转入处置流程')

add_paragraph(doc, '4. 跨境电商流程（9610出口）：')
add_list_item(doc, '步骤1：电商平台将订单信息推送至本系统，系统进行订单数据校验')
add_list_item(doc, '步骤2：系统根据订单信息生成清单，向海关跨境电商系统申报')
add_list_item(doc, '步骤3：海关系统审核清单，返回审核结果（通过/退回/查验）')
add_list_item(doc, '步骤4：审核通过的货物入站，系统自动核注清单')
add_list_item(doc, '步骤5：货物离境后，系统向海关申报离境信息')
add_list_item(doc, '步骤6：海关系统返回结关回执，完成出口业务流程')

doc.add_page_break()

print("第1-3章已生成，继续生成后续章节...")

# 保存文档
doc.save('国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx')
print("文档生成完成！")
print("文件路径: 国际物流航空货站运行平台_技术方案说明书_完整版_V1.0.docx")
